from io import BytesIO

from pypdf import PdfReader
from docx import Document


def extract_pdf_text(file):
    """
    Extract text from an uploaded PDF file.
    """

    reader = PdfReader(file.file)

    text = ""

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text


def extract_docx_text(file):
    """
    Extract text from an uploaded DOCX file.
    """

    document = Document(BytesIO(file.file.read()))

    text = ""

    for para in document.paragraphs:
        if para.text.strip():
            text += para.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text


def extract_resume_text(file):
    """
    Detect the uploaded file type and extract its text.
    """

    filename = file.filename.lower()

    if filename.endswith(".pdf"):
        return extract_pdf_text(file)

    elif filename.endswith(".docx"):
        return extract_docx_text(file)

    raise ValueError("Unsupported file format. Only PDF and DOCX are allowed.")